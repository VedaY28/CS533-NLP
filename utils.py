
import torch
from transformers import AutoTokenizer, AutoModel
from transformers import BertConfig, BertModel
import os
from torch import nn
from torch.utils.data import DataLoader, Dataset
from torch.optim import Adam
from transformers import BertTokenizer, BertModel, get_linear_schedule_with_warmup
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, classification_report
import pandas as pd
import torch.nn.functional as F
from torch import where, rand_like, zeros_like, log, sigmoid
from torch.nn import Module
import matplotlib.pyplot as plt
device = 'cuda' if torch.cuda.is_available() else 'cpu'
from scipy import signal
import numpy as np
from transformers import AutoTokenizer, BertModel
from transformers import  DistilBertModel, DistilBertTokenizerFast
import re
from transformers import AutoTokenizer, AutoModel
from itertools import chain
import regex  # important - not using the usual re module

import docx
from sklearn.feature_extraction.text import CountVectorizer
#from gpu_mem_track import MemTracker
import GPUtil
import gc
from sklearn.metrics.pairwise import cosine_similarity
#from _highlight import highlight_document
import regex as re # important - not using the usual re module
# from sklearn.metrics.pairwise import cosine_similarity
import docx
from scipy.special import softmax
import string
from docx.shared import RGBColor



def check_and_create_folder(path):
    if not os.path.exists(path):
        os.makedirs(path)
        print(f"Created directory: {path}")
    else:
        print(f"Directory already exists: {path}")



import re
alphabets= "([A-Za-z])"
# prefixes = "(Mr|St|Mrs|Ms|Dr)[.]"
suffixes = "(Inc|Ltd|Jr|Sr|Co)"
starters = "(Mr|Mrs|Ms|Dr|Prof|Capt|Cpt|Lt|He\s|She\s|It\s|They\s|Their\s|Our\s|We\s|But\s|However\s|That\s|This\s|Wherever)"
acronyms = "([A-Z][.][A-Z][.](?:[A-Z][.])?)"
websites = "[.](com|net|org|io|gov|edu|me)"
digits = "([0-9])"
multiple_dots = r'\.{2,}'
prefixes = "(Mr|St|Mrs|Ms|Dr|Prof|Capt|Cpt|Lt|Mt)[.]"


def process(text):
    """
    Split the text into sentences.

    If the text contains substrings "<prd>" or "<stop>", they would lead 
    to incorrect splitting because they are used as markers for splitting.

    :param text: text to be split into sentences
    :type text: str

    :return: list of sentences
    :rtype: list[str]
    """
    text = text.lower().split()
    # stops = set(['i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', "you're", "you've", "you'll", "you'd", 'your', 'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', "she's", 'her', 'hers', 'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', "that'll", 'these', 'those', 'am', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while', 'of', 'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', "don't", 'should', "should've", 'now', 'd', 'll', 'm', 'o', 're', 've', 'y'])
    # text = [w for w in text if not w in stops]
    text = " ".join(text)
    text = " " + text + "  "
    text = text.replace("\n"," ")
    text = re.sub(prefixes,"\\1<prd>",text)
    text = re.sub(websites,"<prd>\\1",text)
    text = re.sub(digits + "[.]" + digits,"\\1<prd>\\2",text)
    text = re.sub(multiple_dots, lambda match: "<prd>" * len(match.group(0)) + "<stop>", text)
    if "Ph.D" in text: text = text.replace("Ph.D.","Ph<prd>D<prd>")
    text = re.sub("\s" + alphabets + "[.] "," \\1<prd> ",text)
    text = re.sub(acronyms+" "+starters,"\\1<stop> \\2",text)
    text = re.sub(alphabets + "[.]" + alphabets + "[.]" + alphabets + "[.]","\\1<prd>\\2<prd>\\3<prd>",text)
    text = re.sub(alphabets + "[.]" + alphabets + "[.]","\\1<prd>\\2<prd>",text)
    text = re.sub(" "+suffixes+"[.] "+starters," \\1<stop> \\2",text)
    text = re.sub(" "+suffixes+"[.]"," \\1<prd>",text)
    text = re.sub(" " + alphabets + "[.]"," \\1<prd>",text)
    text = re.sub('<.*?>', ' ', text)
    if "”" in text: text = text.replace("”","")
    if "\"" in text: text = text.replace("\"","")
    # if "!" in text: text = text.replace("!\"","\"!")
    # if "?" in text: text = text.replace("?\"","\"?")
    if "e.g." in text: text = text.replace("e.g.","") 
    if "i.e." in text: text = text.replace("i.e.","")
    if "..." in text: text = text.replace("...","")
    if "<br>" in text: text = text.replace("<br>", "")
    if "<br />" in text: text = text.replace("<br />", "")
    if "(" in text: text = text.replace("(", "")
    if ")" in text: text = text.replace(")", "")
    if "\"\"" in text: text = text.replace("\"\"", "")
    if "\"\'" in text: text = text.replace("\"\'", "")
    if "\"" in text: text = text.replace("\"", "")
    if "*" in text: text = text.replace("*", "")
    if "," in text: text = text.replace(",", " ")
    if "\'" in text: text = text.replace("\'", "")
    if "-" in text: text = text.replace("-", "")
    if "_" in text: text = text.replace("-", "")
    if "<br /><br />" in text: text = text.replace("<br /><br />", "")
    text = text.replace("<prd>",".")
    text = text.replace(".","<stop>")
    text = text.replace("?","<stop>")
    text = text.replace("!","<stop>")
    text = text.replace(";","<stop>")
    text = text.replace(":","<stop>")
    if "<stop>" in text: text = text.replace("<stop>", " ")
    # sentences = text.split("<stop>")
    # sentences = [s.strip() for s in sentences]
    # if sentences and not sentences[-1]: sentences = sentences[:-1]
    # return sentences
    return text.lower().strip()



class TextClassificationDataset_(Dataset):
    def __init__(self, texts):
        self.texts = texts
    def __len__(self):
        return len(self.texts)
    def __getitem__(self, idx):
        text = self.texts[idx]
        return {'original_text': text}

class TextClassificationDataset(Dataset):
    def __init__(self, texts, labels, tokenizer, max_length):
        self.texts = texts
        self.labels = labels
        self.tokenizer = tokenizer
        self.max_length = max_length
        # self.bert = AutoModel.from_pretrained('sentence-transformers/all-mpnet-base-v2')

    def __len__(self):
        return len(self.texts)
    def __getitem__(self, idx):
        text = self.texts[idx]
        processed_text = process(self.texts[idx])
        label = self.labels[idx]
        encoding = self.tokenizer(text, return_tensors='pt', max_length=self.max_length, padding='max_length', truncation=True, is_split_into_words=False, add_special_tokens=True, return_offsets_mapping=True, return_special_tokens_mask=True )
        # embedding = self.bert(input_ids=encoding.input_ids, attention_mask=encoding.attention_mask,  output_hidden_states=True)[0].squeeze()
        return {'original_text': text,  "processed_text": processed_text, 'input_ids': encoding['input_ids'].flatten(), 'attention_mask': encoding['attention_mask'].flatten(), 
                'label': torch.tensor(label), 'offset_mapping': encoding['offset_mapping'].squeeze(), "special_tokens_mask":encoding["special_tokens_mask"].squeeze(),'encoding':encoding}

class ProtoInstanceDataset(Dataset):
    def __init__(self, texts, tokenizer, max_length):
        self.texts = texts
        self.tokenizer = tokenizer
        self.max_length = max_length
    def __len__(self):
        return len(self.texts)
    def __getitem__(self, idx):
        text = self.texts[idx]
        processed_text = process(self.texts[idx])
        encoding = self.tokenizer(self.texts[idx], return_tensors='pt', max_length=self.max_length, padding='max_length', truncation=True, is_split_into_words=False, add_special_tokens=True, return_offsets_mapping=True, return_special_tokens_mask=True )
        return {'original_text': text, "processed_text": processed_text, 'input_ids': encoding['input_ids'].flatten(), 'attention_mask': encoding['attention_mask'].flatten(), 'offset_mapping': encoding['offset_mapping'].squeeze(), "special_tokens_mask":encoding["special_tokens_mask"].squeeze(),'encoding':encoding}

class GoldDataset(Dataset):
    def __init__(self, texts, labels, tokenizer, max_length=512):
        self.texts = texts
        self.labels = labels
        self.tokenizer = tokenizer
        self.max_length = max_length
    def __len__(self):
        return len(self.texts)
    def __getitem__(self, idx):
        text = self.texts[idx]
        label = self.labels[idx]
        processed_text = process(self.texts[idx])
        encoding = self.tokenizer(text, return_tensors='pt', max_length=self.max_length, padding='max_length', truncation=True, is_split_into_words=False, add_special_tokens=True, return_offsets_mapping=True, return_special_tokens_mask=True )
        encoding_label = self.tokenizer(label, return_tensors='pt', max_length=self.max_length, padding='max_length', truncation=True, is_split_into_words=False, add_special_tokens=True, return_offsets_mapping=True, return_special_tokens_mask=True )
        return {'original_text': text, "processed_text": processed_text, 'label_text': label, 
                'input_ids_label': encoding_label['input_ids'].flatten(), 'attention_mask_label': encoding_label['attention_mask'].flatten(), "special_tokens_mask_label":encoding_label["special_tokens_mask"].squeeze().long(),
                'input_ids': encoding['input_ids'].flatten(), 'attention_mask': encoding['attention_mask'].flatten(), 'offset_mapping': encoding['offset_mapping'], "special_tokens_mask":encoding["special_tokens_mask"].squeeze().long(),'encoding':encoding}


def load_data(data_file):
    df = pd.read_csv(data_file)
    texts = df['review'].tolist()
    labels = df['sentiment'].tolist()
    return texts, labels


def load_imdb_data(data_file):
    df = pd.read_csv(data_file)
    texts = df['review'].tolist()
    labels = [1 if sentiment == "positive" else 0 for sentiment in df['sentiment'].tolist()]
    return texts, labels

def get_data_loader(dataset, dataset_path, world_size, rank, batch_size, max_length, bert_model_name, distributed = True):
    data_file = dataset_path
    base = "Datasets/"
    tokenizer = AutoTokenizer.from_pretrained('sentence-transformers/all-mpnet-base-v2')
    if dataset == "IMDB":
        # data_file = dataset_path
        texts, labels = load_imdb_data("/home/bwei2/ProtoTextClassification/Datasets/IMDB/IMDB Dataset.csv")
        train_texts, val_texts, train_labels, val_labels = train_test_split(texts, labels, test_size=0.5, random_state=42)

        # train_texts_df = pd.DataFrame(train_texts, columns=['review'])  # Convert list to DataFrame
        # train_texts_df.to_csv("/home/bwei2/ProtoTextClassification/Data/train_imdb.csv", index=False)
        # test_texts_df = pd.DataFrame(val_texts, columns=['review'])  # Convert list to DataFrame
        # test_texts_df.to_csv("/home/bwei2/ProtoTextClassification/Data/test_imdb.csv", index=False)
    if dataset in ["Hotel", "Amazon", "Rotten_Tomatoes", "Steam", "Yelp", "go_emotion", "Consumer", "ISEAR", "DBPedia"]  :
        train_file = os.path.join(os.path.join(base, dataset),"train.csv")
        test_file = os.path.join(os.path.join(base, dataset),"test.csv")
        
        train_texts,  train_labels = load_data(train_file)
        val_texts, val_labels = load_data(test_file)

    train_dataset = TextClassificationDataset(train_texts, train_labels, tokenizer, max_length)
    val_dataset = TextClassificationDataset(val_texts, val_labels, tokenizer, max_length)
    # train_sampler = torch.utils.data.distributed.DistributedSampler(train_dataset, num_replicas=world_size, rank=rank)
    # train_loader = DataLoader(train_sub, batch_size=batch_size, shuffle=False, pin_memory=pin_memory, sampler=train_sampler)
    train_dataloader = DataLoader(train_dataset, batch_size=batch_size, shuffle=False, pin_memory=True)
    val_dataloader = DataLoader(val_dataset, batch_size=batch_size, shuffle=False, pin_memory=True)
    return train_dataloader, val_dataloader, tokenizer, train_texts



def find_keywords(keywords, text):
    """ Return a list of positions where keywords start or end within the text. 
    Where keywords overlap, combine them. """
    pattern = "(" + "|".join(re.escape(word) for word in keywords) + ")"
    r = []
    for match in regex.finditer(pattern, text, flags=re.I, overlapped=True):
        start, end = match.span()
        if not r or start > r[-1]: 
            r += [start, end]  # add new segment
        elif end > r[-1]:
            r[-1] = end        # combine with previous segment
    return r

def highlight_sections_docx(positions, text, document, label_text=None, key_words=None, matched_proto=None):
    """ Add characters to a text to highlight the segments indicated by
     a list of alternating start and end positions """
    # document = docx.Document()
    # print(positions)
    if label_text is not None:
        p = document.add_paragraph("Prototype sentence:")
        p.add_run(label_text)
    if matched_proto is not None:
        p = document.add_paragraph("Contained prototype sentence:")
        p.add_run(matched_proto)
    if key_words is not None:
        p = document.add_paragraph("Selected chunk: ")
        run = p.add_run(key_words)
        run.bold = True   # or add other formatting - see https://python-docx.readthedocs.io/en/latest/api/text.html#run-objects 

    p = document.add_paragraph("Proto Instance:")
    for i, (start, end) in enumerate(zip([None]+positions , positions+[None])):
        # print("start:", start)
        # print("end:", end)
        run = p.add_run(text[start:end])
        if i % 2:  # odd segments are highlighted
            run.bold = True   # or add other formatting - see https://python-docx.readthedocs.io/en/latest/api/text.html#run-objects 
            run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    return document



import logging
import torch

def longest_common_sublist_torch(list1, list2):

    dp = torch.zeros(len(list1) + 1, len(list2) + 1, dtype=torch.int)

    longest_length = 0
    longest_end_pos = 0

    # Populate the DP table using torch.eq for element-wise comparison
    for i in range(1, len(list1) + 1):
        for j in range(1, len(list2) + 1):
            if torch.eq(list1[i - 1], list2[j - 1]):
                dp[i][j] = dp[i - 1][j - 1] + 1
                if dp[i][j] > longest_length:
                    longest_length = dp[i][j]
                    longest_end_pos = i

    # Extract the start and end of the longest common sublist
    start = longest_end_pos - longest_length
    end = longest_end_pos
    # logging.info(list1[start:end])
    # # Create a tensor filled with zeros
    # result = torch.zeros_like(list1, dtype=torch.int)
    
    # # Use torch.where to set the elements in the range [start:end] to 1
    # result = torch.where((torch.arange(len(list1)) >= start) & (torch.arange(len(list1)) < end), 1, result)

    return list(range(start, end + 1))

def get_selected_token_index(label_mask, start_point_ids, num_prototypes):
    # Get the batch size
    batch_size = label_mask.shape[0]# // num_prototypes
    all_selected_token_index = []
    for i in range(batch_size):
        # print("i",i)
        for j in range(num_prototypes):
            label_mask_list = list(label_mask[i, j].cpu().numpy())
            target = start_point_ids[i*num_prototypes+ j]
            flag = True
            key=False
            for token_num,token_id in enumerate(target):
                
                if token_id in label_mask_list:
                    temp = label_mask_list.index(token_id)
                    # if token_num+1 < len(target) and label_mask_list[temp+1] == target[token_num+1]:
                    flag = False
                    all_selected_token_index.append(temp)#longest_common_sublist_torch(label_mask[i, j], start_point_ids[i*num_prototypes+ j]))
                    break
            if flag  :
                all_selected_token_index.append(0)
    # print(len(all_selected_token_index))
    return all_selected_token_index # Reshape back to original dimensions






def test_intermediate_results(test_text_list, tokenizer, model, device): 

    for idx, test_text in enumerate(test_text_list): 
        logging.info("test text {}: {} \n".format(idx, test_text))
        encoding = tokenizer(test_text, return_tensors='pt', max_length=model.max_length, padding='max_length', truncation=True)
        input_ids_test = encoding['input_ids'].to(device)
        ids_input = encoding['input_ids']
        test_ori = tokenizer.batch_decode(ids_input)
        attention_mask_test = encoding['attention_mask'].to(device)
        outputs_test = model.bert(input_ids=input_ids_test, attention_mask=attention_mask_test)
        hidden_states = outputs_test.last_hidden_state
        mask = model.attention_layer(hidden_states, model.prototype_vectors)
        input_ids_new = input_ids_test.unsqueeze(-1) * mask
        # input_ori = tokenizer.decode(input_ids_new[:, :, 0])
        output_sub_seq_emb = torch.cat([model.bert(input_ids= input_ids_new[:, :, i].long(), attention_mask=attention_mask_test).last_hidden_state[:, 0, :].unsqueeze(-1) for i in range(model.num_prototypes)], 2)
        z_prime = output_sub_seq_emb.view(-1, model.num_prototypes, model.hidden_dim)
        cos = nn.CosineSimilarity(dim=2, eps=1e-6)
        similarity = F.softmax(cos(z_prime, model.prototype_vectors), dim=1)

        for proto_num in range(model.num_prototypes):
            ids_masked = input_ids_test * mask[0,:, proto_num]
            test_masked = tokenizer.batch_decode(ids_masked.int())
            TEST_UNTOKENized = tokenizer.convert_tokens_to_string(test_masked).replace("[PAD]", "")
            
            new_masked_text = list(map(lambda x: x.replace('[PAD]','_'),test_masked))
            # logging.info(' '.join([str(elem) for elem in new_masked_text]))
            logging.info(f"Similarity: {similarity[:,proto_num].item():.4f}\t\n Prototype {proto_num}: {TEST_UNTOKENized}" )
    

def remove_deactivated_elements(pi, mu, sigma, threshold=0.15):
    active_mask = pi > threshold  # Shape: (batch_size, n_components)

    has_active = active_mask.any(dim=1)  # Shape: (batch_size,)
    
    argmax_indices = pi.argmax(dim=1)  # Shape: (batch_size,)
    

    final_active_pi_list = []
    final_active_mu_list = []
    final_active_sigma_list = []
    

    final_active_pi_list = [
        pi[i][active_mask[i]] if has_active[i] else pi[i, argmax_indices[i]].unsqueeze(0)
        for i in range(pi.size(0))
    ]
    
    final_active_mu_list = [
        mu[i][active_mask[i]] if has_active[i] else mu[i, argmax_indices[i]].unsqueeze(0)
        for i in range(mu.size(0))
    ]
    
    final_active_sigma_list = [
        sigma[i][active_mask[i]] if has_active[i] else sigma[i, argmax_indices[i]].unsqueeze(0)
        for i in range(sigma.size(0))
    ]
    
    return final_active_pi_list, final_active_mu_list, final_active_sigma_list
